"""
Document Parser Factory - 支持多种文档类型的解析器
"""
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
import os
import re


class BaseParser(ABC):
    """解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文档，返回 {'text': str, 'metadata': dict}"""
        pass

    @abstractmethod
    def get_type(self) -> str:
        """返回解析器类型"""
        pass

    def _fallback_parse(self, file_path: str, reason: str = "") -> Dict[str, Any]:
        """解析失败时的后备方案

        `reason` is the exception/error string from the original parser.
        Downstream code (document_tasks, knowledge upload) reads
        ``metadata.parse_error`` to decide whether to mark the
        ``Document`` row as failed instead of silently committing
        garbage chunks.
        """
        text = self._read_text(file_path)
        meta = {"type": self.get_type(), "fallback": True}
        if reason:
            meta["parse_error"] = reason
        return {
            "text": text,
            "metadata": meta,
        }

    def _read_text(self, file_path: str) -> str:
        """读取文本文件"""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
        except OSError:
            # File missing or unreadable (e.g. docling deleted the
            # temp file before falling back). Return empty — the
            # caller still records the original parse error so the
            # Document row ends up in the right state.
            return ""
        return ""

    def _looks_like_pdf_byte_stream(self, text: str) -> bool:
        """检测 Docling 输出的"文本"是否实际上是 PDF 字节流(乱码)。

        Docling 在字体 ToUnicode CMap 缺失的 PDF 上(典型如 Chrome headless
        打印的 PDF)会把 PDF 字节流本身当作"提取出的文本"返回。返回的
        文本仍然非空、``parse_error`` 仍然为 None,所以下游的
        ``if not text_content or parse_error`` 守门会漏掉它,导致乱码
        继续走 chunking + embedding,最终入库成 ``vector_id=error_N`` 的
        乱码 chunks 且 ``doc.status=completed``(数据三重不一致)。

        启发式:任一命中即返回 True。
          1. 文本以 ``%PDF-`` 开头
          2. 前 2000 字符含 2+ PDF 内部对象标记
             (FlateDecode / endobj / endstream / MediaBox / /Type /Page / /Encoding)
          3. 前 2000 字符内不可打印控制字符(不含 ``\\n\\r\\t``)比例 > 5%

        空文本与短文本(< 50 字符)直接返回 False,避免误杀标题或
        "no extractable text" 的极短返回。
        """
        if not text:
            return False

        sample = text[:2000]

        # Rule 1: PDF header (strong signal — fires even on short text)
        if sample.lstrip().startswith("%PDF-"):
            return True

        # Below this length, marker-counting / control-char-ratio signals
        # are too noisy (e.g. "FlateDecode endobj endstream" could be a
        # paper title). The PDF header above is the only short-circuit.
        if len(text) < 50:
            return False

        # Rule 2: 2+ PDF internal markers in sample window
        markers = ("FlateDecode", "endobj", "endstream", "MediaBox", "/Type /Page", "/Encoding")
        marker_hits = sum(sample.count(m) for m in markers)
        if marker_hits >= 2:
            return True

        # Rule 3: > 5% non-printable control chars (excluding \n\r\t)
        if not sample:
            return False
        non_printable = sum(
            1 for c in sample
            if ord(c) < 0x20 and c not in "\n\r\t"
        )
        if non_printable / len(sample) > 0.05:
            return True

        return False

    def _parse_with_pdfplumber(self, file_path: str) -> str:
        """Level 2 fallback: pdfplumber text extraction.

        Returns the joined text of all pages, or raises on failure.
        Caller is responsible for the empty / garbage-text check.
        """
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)

    def _parse_with_pypdfium2(self, file_path: str) -> str:
        """Level 3 fallback: pypdfium2 text extraction.

        Returns the joined text of all pages, or raises on failure.
        Caller is responsible for the empty / garbage-text check.
        """
        import pypdfium2 as pdfium  # type: ignore[import-untyped]
        pdf = pdfium.PdfDocument(file_path)
        try:
            chunks = []
            for i in range(len(pdf)):
                text_page = pdf[i].get_textpage()
                try:
                    chunks.append(text_page.get_text_range())
                finally:
                    text_page.close()
            return "\n".join(chunks)
        finally:
            pdf.close()

    def _parse_with_fallback(
        self, file_path: str, primary_parser_func
    ) -> Dict[str, Any]:
        """统一 fallback 入口 (M29.2 扩展):按文件扩展名分发到 PDF 或 docx chain。

        ``primary_parser_func(file_path)`` 走第一个解析器(默认是 Docling)。
        - ``.docx`` → ``_parse_with_docx_fallback``(Docling → python-docx → _fallback_parse)
        - 其他(默认 PDF)→  ``_parse_with_pdf_fallback``(Docling → pdfplumber → pypdfium2 → _fallback_parse)

        所有 6 个 parser (General/Paper/QA/Table/Manual/Laws) 都调本入口,
        改 fallback 行为不需要分别改 6 个 ``parse()``。
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return self._parse_with_docx_fallback(file_path, primary_parser_func)
        return self._parse_with_pdf_fallback(file_path, primary_parser_func)

    def _parse_with_pdf_fallback(
        self, file_path: str, primary_parser_func
    ) -> Dict[str, Any]:
        """三级 PDF 解析链:primary → pdfplumber → pypdfium2 → _fallback_parse。

        ``primary_parser_func(file_path)`` 走第一个解析器(默认是 Docling)。
        任一级返回的 text 如果是 PDF 字节流乱码(用 ``_looks_like_pdf_byte_stream``
        检测)或抛异常,就降级到下一级。三级都失败时调 ``_fallback_parse``,
        ``metadata.parse_error`` 会写明最后一级失败的原因 —— 下游的
        ``document_tasks`` / ``knowledge.py`` 据此把 ``Document.status`` 设为
        FAILED,而不是静默提交空/乱码 chunks。
        """
        last_error = ""

        # Level 1: primary (typically Docling)
        try:
            result = primary_parser_func(file_path)
            if not self._looks_like_pdf_byte_stream(result.get("text", "")):
                return result
            last_error = "primary parser returned garbage (looks like PDF byte stream)"
        except Exception as e:
            last_error = f"primary: {type(e).__name__}: {e}"

        # Level 2: pdfplumber
        try:
            text = self._parse_with_pdfplumber(file_path)
            if text and not self._looks_like_pdf_byte_stream(text):
                return {
                    "text": text,
                    "metadata": {
                        "type": self.get_type(),
                        "parser": "pdfplumber_fallback",
                    },
                }
            if not text:
                last_error = "pdfplumber returned empty text"
        except Exception as e:
            last_error = f"pdfplumber: {type(e).__name__}: {e}"

        # Level 3: pypdfium2
        try:
            text = self._parse_with_pypdfium2(file_path)
            if text and not self._looks_like_pdf_byte_stream(text):
                return {
                    "text": text,
                    "metadata": {
                        "type": self.get_type(),
                        "parser": "pypdfium2_fallback",
                    },
                }
            if not text:
                last_error = "pypdfium2 returned empty text"
        except Exception as e:
            last_error = f"pypdfium2: {type(e).__name__}: {e}"

        # All three levels failed — record the last error so the
        # downstream Document row ends up in the right state.
        return self._fallback_parse(
            file_path,
            reason=last_error or "all PDF parsers failed",
        )

    # ----- DOCX 三级 fallback chain (M29.2, 2026-06-15) -----------------------
    #
    # 背景: Docling 2.24 对 Word ``numPr``(列表样式)段处理有结构性 bug —
    # 把"容器项"(``车辆管理``)与后续 ``页面展示:`` 段合并,导致 list-style
    # 的"页面展示"内容被吞(精准停车功能文档.docx 实测:Docling 输出
    # 15,448 字符丢 36%,``页面展示`` 段从 32/32 降到 12/32)。
    #
    # 修法: 仿 PDF 的 pdfplumber + pypdfium2 三级 chain,加
    # Docling → python-docx → _fallback_parse 链;守门启发式:
    # Docling 输出字符数 < OOXML ``word/document.xml`` 文本长度 × 0.7 时
    # 触发 fallback(精准停车文档实测 0.62,会触发)。

    def _docx_xml_text_length(self, file_path: str) -> int:
        """读 OOXML ``word/document.xml``,sum 所有 ``<w:t>`` 文本长度。

        作为 docx 截断守门的"原文字符数"基准。docx 是 zip 格式,
        不会触发文件级 I/O 错误,除 zipfile.BadZipFile / KeyError
        (非 docx 文件)外都返回整数。
        """
        import zipfile
        try:
            with zipfile.ZipFile(file_path) as z:
                with z.open("word/document.xml") as f:
                    xml_bytes = f.read()
        except (zipfile.BadZipFile, KeyError, OSError):
            return 0

        # 抽所有 <w:t>...</w:t> 标签内的文本,用 regex 累加长度。
        # 用 namespace 不强制 — Word 输出始终是
        # ``<w:t xml:space="preserve">…</w:t>``,非贪婪即可。
        import re as _re
        matches = _re.findall(rb"<w:t[^>]*>([^<]*)</w:t>", xml_bytes)
        return sum(len(m.decode("utf-8", errors="replace")) for m in matches)

    def _parse_with_python_docx(self, file_path: str) -> str:
        """Level 2 fallback: 直接读 OOXML 用 python-docx。

        模拟 Word 列表样式渲染: ``numPr`` 段加 ``- `` 前缀,与 Docling
        输出大致对齐(让 chunking_service 一致处理)。python-docx 不
        解析嵌入图片/复杂表格 — 当前 KB 文档以文字为主,RAG 弱场景。
        """
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        # OOXML namespace prefix,python-docx 内部用
        NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            # numPr 段: list-style 列表项,加 ``- `` 前缀
            if p._element.find(f".//{NS}numPr") is not None:
                paragraphs.append(f"- {text}")
            else:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)

    def _looks_like_truncated_docx(self, text: str, file_path: str) -> bool:
        """守门: Docling 输出字符数 < OOXML ``<w:t>`` 总长度 × 0.7 视为截断。

        实测精准停车功能文档.docx: OOXML ``<w:t>`` 24,872 chars,
        Docling 输出 15,448 chars (62%),触发 fallback。
        阈值保守设 0.7 — 若未来健康 Docling 输出 < 70% (unlikely)
        误判,只会走 Level 2 python-docx,不影响主流程。
        """
        if not text:
            return False
        ooxml_len = self._docx_xml_text_length(file_path)
        if ooxml_len <= 0:
            return False  # 非 docx 文件(可能 PDF 走同一 chain)放过
        return len(text) < ooxml_len * 0.7

    def _parse_with_docx_fallback(
        self, file_path: str, primary_parser_func
    ) -> Dict[str, Any]:
        """三级 docx 解析链:primary → python-docx → _fallback_parse。

        与 PDF chain 平行实现,``primary_parser_func(file_path)`` 走
        Docling,触发守门 ``_looks_like_truncated_docx`` 时降级到
        python-docx。docx 路径不调 PDF 专用的 ``_looks_like_pdf_byte_stream``
        (PDF 路径独立在 ``_parse_with_fallback`` 里)。
        """
        last_error = ""

        # Level 1: primary (typically Docling)
        try:
            result = primary_parser_func(file_path)
            text = result.get("text", "")
            if not self._looks_like_truncated_docx(text, file_path):
                return result
            last_error = (
                f"primary parser returned truncated docx "
                f"(text len {len(text)} < 70% of OOXML)"
            )
        except Exception as e:
            last_error = f"primary: {type(e).__name__}: {e}"

        # Level 2: python-docx
        try:
            text = self._parse_with_python_docx(file_path)
            if text:
                return {
                    "text": text,
                    "metadata": {
                        "type": self.get_type(),
                        "parser": "python_docx_fallback",
                        "fallback_reason": last_error,
                    },
                }
            last_error = "python-docx returned empty text"
        except Exception as e:
            last_error = f"python-docx: {type(e).__name__}: {e}"

        # Both levels failed — record the last error
        return self._fallback_parse(
            file_path,
            reason=last_error or "all docx parsers failed",
        )


class GeneralParser(BaseParser):
    """通用解析器 - 基础文档类型"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """使用 Docling 解析通用文档,失败/乱码时降级到 pdfplumber/pypdfium2"""
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            return {
                "text": result.document.export_to_text(),
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "general"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""


class PaperParser(BaseParser):
    """学术论文解析器 - 优化处理论文结构"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            text = result.document.export_to_text()
            return {
                "text": text,
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                    "authors": self._extract_authors(result),
                    "sections": self._extract_sections(text),
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "paper"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""

    def _extract_authors(self, result) -> List[str]:
        try:
            authors = result.document.metadata.get("authors", [])
            return [a.get("name", "") for a in authors] if isinstance(authors, list) else []
        except:
            return []

    def _extract_sections(self, text: str) -> List[Dict[str, str]]:
        sections = []
        patterns = [
            r"(?:^|\n)([一二三四五六七八九十]+[、.]\s*[^\n]+)",
            r"(?:^|\n)(Abstract|Introduction|Related Work|Methodology|Experiment|Conclusion|References)",
            r"(?:^|\n)(第[一二三四五六七八九十]+[章节])",
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for m in matches:
                title = m.group(1).strip()
                if title and len(title) < 100:
                    sections.append({"title": title, "position": m.start()})
        return sections[:20]


class QAParser(BaseParser):
    """问答解析器 - 从文档中提取问答对"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            text = result.document.export_to_text()
            qa_pairs = self._extract_qa_pairs(text)
            return {
                "text": text,
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                    "qa_pairs": qa_pairs,
                    "count": len(qa_pairs),
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "qa"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""

    def _extract_qa_pairs(self, text: str) -> List[Dict[str, str]]:
        qa_pairs = []
        qa_patterns = [
            r"问[：:]\s*([^\n问]+)\s*答[：:]\s*([^\n]+)",
            r"(?:^|\n)(Q[：:]\s*[^\n]+)\s*(?:A[：:]\s*[^\n]+)",
            r"题目[：:]\s*([^\n]+)\s*答案[：:]\s*([^\n]+)",
            r"(?:^|\n)(\d+[.、)]\s*[^\n?]+[?？])\s*\n?\s*([^\n]+?)(?=\n\d+[.、)]\s*|$)",
        ]

        for pattern in qa_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for m in matches:
                if len(m.groups()) >= 2:
                    q = m.group(1).strip()
                    a = m.group(2).strip()
                    if q and a and len(q) > 2 and len(a) > 2:
                        qa_pairs.append({"question": q, "answer": a})

        seen = set()
        unique_pairs = []
        for pair in qa_pairs:
            key = pair["question"][:50]
            if key not in seen:
                seen.add(key)
                unique_pairs.append(pair)

        return unique_pairs[:100]


class TableParser(BaseParser):
    """表格解析器 - 保留表格结构"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            text = result.document.export_to_text()
            tables = []

            try:
                for item in result.document.pages:
                    if hasattr(item, 'tables'):
                        for table in item.tables:
                            tables.append(self._table_to_markdown(table))
            except Exception:
                pass

            return {
                "text": text,
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                    "tables": tables,
                    "table_count": len(tables),
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "table"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""

    def _table_to_markdown(self, table) -> str:
        try:
            return table.export_to_markdown()
        except:
            return ""


class ManualParser(BaseParser):
    """用户手册解析器 - 优化处理手册结构"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            text = result.document.export_to_text()
            steps = self._extract_steps(text)
            warnings = self._extract_warnings(text)
            return {
                "text": text,
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                    "steps": steps,
                    "warnings": warnings,
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "manual"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""

    def _extract_steps(self, text: str) -> List[str]:
        steps = []
        patterns = [
            r"(?:^|\n)(\d+[.、)]\s*(?:第[一二三四五六七八九十]+步[：:]?\s*)?[^\n]+)",
            r"(?:^|\n)(步骤[一二三四五六七八九十]+[：:]?\s*[^\n]+)",
            r"(?:^|\n)([首先然后最后此外]+[，,\s]+[^\n]+)",
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                step = m.group(1).strip()
                if step and len(step) < 200:
                    steps.append(step)
        return steps[:50]

    def _extract_warnings(self, text: str) -> List[str]:
        warnings = []
        patterns = [
            r"(?:警告|注意|注意啦| Caution | Warning )[：:]\s*([^\n]+)",
            r"(?:重要|请勿|不要)[：:]\s*([^\n]+)",
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for m in matches:
                warning = m.group(1).strip()
                if warning:
                    warnings.append(warning)
        return warnings[:20]


class LawsParser(BaseParser):
    """法律文档解析器 - 保留条款结构"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        def docling_parse(fp: str) -> Dict[str, Any]:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(fp)
            text = result.document.export_to_text()
            articles = self._extract_articles(text)
            return {
                "text": text,
                "metadata": {
                    "type": self.get_type(),
                    "title": self._extract_title(result),
                    "articles": articles,
                    "article_count": len(articles),
                }
            }
        return self._parse_with_fallback(file_path, docling_parse)

    def get_type(self) -> str:
        return "laws"

    def _extract_title(self, result) -> str:
        try:
            return result.document.metadata.get("title", "")
        except:
            return ""

    def _extract_articles(self, text: str) -> List[Dict[str, Any]]:
        articles = []
        patterns = [
            r"(?:第[一二三四五六七八九十百]+[条章节])[：:\s]*([^\n]+)",
            r"(?:^|\n)(\d+[.。]\s*[^\n]+?(?:条|款|项)[^\n]*)",
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                article = m.group(1).strip()
                if article and len(article) < 500:
                    articles.append({"text": article, "position": m.start()})
        return articles[:100]


class DocumentParserFactory:
    """文档解析器工厂"""

    PARSERS = {
        "general": GeneralParser,
        "paper": PaperParser,
        "qa": QAParser,
        "table": TableParser,
        "manual": ManualParser,
        "laws": LawsParser,
    }

    # 文件名模式到解析器类型的映射
    TYPE_PATTERNS = {
        "paper": [
            r"paper", r" Paper", r" research", r"学术", r"论文",
            r"conference", r"journal", r"arxiv"
        ],
        "qa": [
            r"qa", r" Q&A", r"question", r"answer", r"问答",
            r"faq", r"FAQ", r"面试题", r"题库"
        ],
        "table": [
            r"table", r" Table", r"表格", r"数据表",
            r"excel", r"spreadsheet"
        ],
        "manual": [
            r"manual", r" Manual", r"guide", r" Guide",
            r"手册", r"指南", r"说明书", r"user guide",
            # M29.2 (2026-06-15): 功能文档/需求文档/产品文档/PRD 命中 → ManualParser
            # → document_structure 分块策略(替换原 general → fixed_size_500_50),
            # 关键修复 5 个 ``（一）`` 二级标题孤立成 21 字符 chunk 的问题。
            r"功能文档", r"需求文档", r"产品文档", r"PRD",
        ],
        "laws": [
            r"law", r" Laws", r"legal", r"contract", r"agreement",
            r"法律", r"法规", r"条款", r"合同", r"协议"
        ],
    }

    # 内容特征模式
    CONTENT_PATTERNS = {
        "paper": {
            "keywords": [
                r"abstract", r"introduction", r"related work", r"methodology",
                r"experiment", r"conclusion", r"references", r"bibliography",
                r"摘要", r"引言", r"相关工作", r"实验", r"结论", r"参考文献"
            ],
            "weight": 1.0
        },
        "qa": {
            "keywords": [
                r"问[：:]", r"答[：:]", r"Q[：:]\s*A",
                r"问题[：:]", r"答案[：:]",
                r"^\d+[.、]\s*[^\n?]+[?？]", r"^\d+[.、]\s*答案",
                r"面试题", r"题库", r"考题"
            ],
            "weight": 1.2
        },
        "table": {
            "keywords": [
                r"\|.+\|", r"\|.+\|.*\|",
                r"^\s*[\d\.]+\s+\S+\s+\S+",
                r"表格", r"数据", r"统计",
                r"月份", r"金额", r"数量"
            ],
            "weight": 1.0
        },
        "manual": {
            "keywords": [
                r"步骤", r"操作指南", r"使用方法", r"注意事项",
                r"warning", r"caution", r"notice",
                r"警告", r"注意", r"请勿", r"important",
                r"第\d+步", r"步骤\d+", r"操作步骤"
            ],
            "weight": 1.0
        },
        "laws": {
            "keywords": [
                r"第[一二三四五六七八九十百]+条",
                r"当事人", r"甲方", r"乙方", r"违约",
                r"条款", r"协议", r"合同法",
                r"article\s+\d+", r"section\s+\d+",
                r"甲方", r"乙方", r"签署", r"生效"
            ],
            "weight": 1.0
        },
    }

    @classmethod
    def get_parser(cls, doc_type: str = None, file_path: str = None) -> BaseParser:
        """获取解析器"""
        if doc_type and doc_type in cls.PARSERS:
            return cls.PARSERS[doc_type]()

        if file_path:
            detected_type = cls.detect_doc_type(file_path)
            if detected_type:
                return cls.PARSERS[detected_type]()

        return GeneralParser()

    @classmethod
    def detect_doc_type(cls, file_path: str) -> Optional[str]:
        """根据文件名检测文档类型"""
        filename = os.path.basename(file_path).lower()

        for doc_type, patterns in cls.TYPE_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, filename, re.IGNORECASE):
                    return doc_type

        return None

    @classmethod
    def detect_doc_type_from_content(cls, text: str, min_score: float = 0.5) -> Optional[str]:
        """根据内容检测文档类型"""
        if not text:
            return None

        sample_text = text[:2000].lower()

        scores = {}
        for doc_type, config in cls.CONTENT_PATTERNS.items():
            score = 0.0
            keyword_count = 0

            for pattern in config["keywords"]:
                matches = len(re.findall(pattern, sample_text, re.IGNORECASE))
                if matches > 0:
                    keyword_count += matches
                    score += matches * config["weight"]

            if keyword_count > 0:
                scores[doc_type] = score / max(1, len(config["keywords"]))

        if not scores:
            return None

        best_type = max(scores.items(), key=lambda x: x[1])

        if best_type[1] >= min_score:
            return best_type[0]

        return None

    @classmethod
    def auto_detect(
        cls,
        file_path: str = None,
        text_sample: str = None,
        use_content: bool = True
    ) -> str:
        """自动检测文档类型"""
        # 1. 先尝试文件名检测（快速）
        if file_path:
            filename_type = cls.detect_doc_type(file_path)
            if filename_type:
                return filename_type

        # 2. 再尝试内容检测（更准确但需要文本）
        if use_content and text_sample:
            content_type = cls.detect_doc_type_from_content(text_sample)
            if content_type:
                return content_type

        # 3. 默认通用类型
        return "general"

    @classmethod
    def get_available_types(cls) -> List[Dict[str, str]]:
        """获取可用的文档类型"""
        return [
            {"type": "general", "label": "通用文档", "description": "通用解析器，支持 PDF/DOCX/TXT 等"},
            {"type": "paper", "label": "学术论文", "description": "优化处理论文结构，提取作者、章节"},
            {"type": "qa", "label": "问答文档", "description": "从文档中提取问答对"},
            {"type": "table", "label": "表格文档", "description": "保留表格结构，提取表格内容"},
            {"type": "manual", "label": "用户手册", "description": "提取操作步骤和警告信息"},
            {"type": "laws", "label": "法律文档", "description": "保留条款结构，提取法律条款"},
        ]
