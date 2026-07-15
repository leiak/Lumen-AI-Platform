from typing import List, Dict, Any
import io

class DocumentGenerator:
    """Generate documents in various formats"""

    def generate_word(self, title: str, content: str) -> bytes:
        """Generate a Word document"""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(title, 0)

        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_excel(self, data: List[Dict[str, Any]], headers: List[str] = None) -> bytes:
        """Generate an Excel spreadsheet"""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active

        if not data:
            return io.BytesIO().getvalue()

        # Headers
        if headers is None:
            headers = list(data[0].keys()) if data else []

        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)

        # Data rows
        for row, item in enumerate(data, 2):
            for col, header in enumerate(headers, 1):
                ws.cell(row=row, column=col, value=item.get(header, ""))

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()